#!/usr/bin/env python
# -*- coding:utf-8 -*-
# __author__ = hzmin
# __date__ = 2022/1/6

import datetime

from django.conf import settings
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.viewsets import ViewSet

from utils.fake import fake
from utils.pdf import html2pdf, html2doc


class DocViewSet(ViewSet):
    authentication_classes = ()
    permission_classes = ()

    @action(methods=["GET"], detail=False)
    def doc(self, request):
        context = {
            "logo_url": f"{settings.BASE_DIR}/ignore_data/4k.jpg",  # 保险公司logo图片
            "seal_url": f"{settings.BASE_DIR}/ignore_data/zhang.png",  # 购买方章
            "company_name": "company_name",
            "insurance_detail_name": f"xxx-xxx",
            "policy_no": "policy_no",
            "date": datetime.datetime.now(),
            "headers": [
                "序号",
                "姓名",
                "证件类型",
                "证件号码",
                "工种",
                "保险起期\n(日期固定格式)",
                "保险止期\n(日期固定格式)",
                "职业类别",
                "增员/减员",
            ],
            "results": [[i + 1, fake.name(), "身份证", fake.ssn(), fake.word(), "2022-01-01", "2022-01-31", 5, "增员"]
                        for i in range(20)],
        }

        template_name = f"eli_import_list_2_sent"
        pdf_data = html2pdf(template_name, context)

        pdf_file = f"{settings.BASE_DIR}/ignore_data/aaa1.pdf"
        with open(pdf_file, "wb") as f:
            f.write(pdf_data)

        docx_file = f"{settings.BASE_DIR}/ignore_data/aaa3.docx"
        html2doc(pdf_file, docx_file)

        import docx

        def del_table(table):
            # 适用于删除段落、表格、图片等
            t = table._element
            t.getparent().remove(t)
            t._p = t._element = None

        def move_table_after(table, doc, s):
            # 在指定文本后插入表格
            prev = None
            for pg in doc.paragraphs:
                if s in pg.text:
                    prev = pg
            if prev:
                t, p = table._tbl, prev._p
                p.addnext(t)

        doc = docx.Document(docx_file)

        import copy
        old_r = None
        for table in doc.tables:
            if len(table.rows[0].cells) != 9:
                del_table(table)
            else:
                if old_r is None:
                    old_r = copy.deepcopy(table.rows[0])
                table.rows[0]._tr.addprevious(copy.deepcopy(old_r)._element)
                for c, v in zip(table.rows[0].cells, context["headers"]):
                    c.text = v
                table.style = "Table Grid"
                table.autofit = True
                for r in table.rows:
                    from docx.enum.table import WD_ROW_HEIGHT_RULE
                    r.height_rule = WD_ROW_HEIGHT_RULE.AUTO

        doc.save(f"{settings.BASE_DIR}/ignore_data/aaa4.docx")
        return Response("ok")
